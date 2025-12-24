from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r'Roommate Agreement Generator – Design & Architecture (azure + Python).docx')
for para in doc.paragraphs:
    print(para.text)
    
# Also extract tables if any
for table in doc.tables:
    print("\n--- TABLE ---")
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(" | ".join(row_text))
